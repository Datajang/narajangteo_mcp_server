#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Universal Content Extraction Module
다형식 파일 처리 모듈 - HWP, HWPX, PDF, DOCX, XLSX, ZIP
"""

import io
import os
import re
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional

import httpx

# Optional imports with fallbacks
try:
    import olefile
    HAS_OLEFILE = True
except ImportError:
    HAS_OLEFILE = False

try:
    from langchain_teddynote.document_loaders import HWPLoader
    HAS_HWPLOADER = True
except ImportError:
    HAS_HWPLOADER = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


def select_best_file_from_zip(file_list: list[str]) -> Optional[str]:
    """
    ZIP 파일 내에서 가장 적합한 파일 선택 (가이드 §8.2.A 준수)

    우선순위:
    1. 파일명에 '제안요청서' 또는 '과업지시서' 포함
    2. 확장자 .hwp 또는 .hwpx
    3. 확장자 .docx 또는 .pdf
    4. 없으면 None
    """
    # 숨김 파일 및 __MACOSX 제외
    valid_files = [
        f for f in file_list
        if not f.startswith('__MACOSX')
        and not os.path.basename(f).startswith('.')
        and not f.endswith('/')
    ]

    # Priority 1: 제안요청서/과업지시서
    for f in valid_files:
        basename = os.path.basename(f)
        if '제안요청서' in basename or '과업지시서' in basename:
            return f

    # Priority 2: HWP/HWPX
    for f in valid_files:
        if f.lower().endswith(('.hwp', '.hwpx')):
            return f

    # Priority 3: DOCX/PDF
    for f in valid_files:
        if f.lower().endswith(('.docx', '.pdf')):
            return f

    return None


def _extract_from_hwp_olefile(file_bytes: bytes) -> str:
    """
    Fallback: olefile을 사용한 HWP 추출 (기존 구현)
    DRM/암호화 시 "HWP Protected" 반환
    """
    if not HAS_OLEFILE:
        return "HWP extraction requires olefile or langchain-teddynote library."

    try:
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))

        # Check for encryption
        if ole.exists('EncryptedPackage'):
            return "HWP Protected: This file is encrypted."

        # Try to extract text from various streams
        text_parts = []

        # PrvText stream contains preview text
        if ole.exists('PrvText'):
            try:
                prv_text = ole.openstream('PrvText').read()
                # PrvText is UTF-16LE encoded
                decoded = prv_text.decode('utf-16-le', errors='ignore')
                # Remove null characters
                decoded = decoded.replace('\x00', '')
                if decoded.strip():
                    text_parts.append(decoded)
            except Exception:
                pass

        # BodyText sections
        for i in range(100):
            section_name = f'BodyText/Section{i}'
            if ole.exists(section_name):
                try:
                    section_data = ole.openstream(section_name).read()
                    # Try to decode as UTF-16LE
                    try:
                        decoded = section_data.decode('utf-16-le', errors='ignore')
                        # Extract readable text using regex
                        readable = re.findall(r'[\uAC00-\uD7A3a-zA-Z0-9\s.,!?@#$%^&*()_+=\-\[\]{}|;:\'\"<>/\\]+', decoded)
                        text_parts.extend(readable)
                    except Exception:
                        pass
                except Exception:
                    pass
            else:
                break

        ole.close()

        if text_parts:
            return '\n'.join(text_parts)
        else:
            return "HWP: Could not extract text. File may use unsupported encoding."

    except Exception as e:
        return f"HWP extraction failed: {str(e)}"


def extract_from_hwp(file_bytes: bytes) -> str:
    """
    HWP 파일에서 텍스트 추출
    1차: HWPLoader (langchain-teddynote) - 높은 정확도, zlib 압축 지원
    2차: olefile 직접 파싱 - 폴백
    """
    # Primary: Try HWPLoader
    if HAS_HWPLOADER:
        temp_file_path = None
        try:
            # 1. 임시 파일 생성 (.hwp 확장자 필수)
            with tempfile.NamedTemporaryFile(
                mode='wb',
                suffix='.hwp',
                delete=False
            ) as temp_file:
                temp_file.write(file_bytes)
                temp_file_path = temp_file.name

            # 2. HWPLoader로 로드
            loader = HWPLoader(file_path=temp_file_path)
            documents = list(loader.lazy_load())

            # 3. Document에서 텍스트 추출
            if documents:
                text_parts = [
                    doc.page_content
                    for doc in documents
                    if doc.page_content.strip()
                ]
                if text_parts:
                    return '\n\n'.join(text_parts)

            # HWPLoader 성공했지만 텍스트 없음 → olefile fallback
            return _extract_from_hwp_olefile(file_bytes)

        except Exception as e:
            # HWPLoader 실패 → olefile fallback
            # 디버깅을 위한 로그 출력
            import sys
            print(f"[DEBUG] HWPLoader failed, falling back to olefile: {str(e)}", file=sys.stderr)
            return _extract_from_hwp_olefile(file_bytes)

        finally:
            # 4. 임시 파일 정리
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except:
                    pass  # 정리 실패 무시

    # HWPLoader 없음 → olefile 직접 사용
    else:
        return _extract_from_hwp_olefile(file_bytes)


def extract_from_hwpx(file_bytes: bytes) -> str:
    """
    HWPX 파일에서 텍스트 추출 (ZIP + XML 파싱)
    Contents/section0.xml에서 텍스트 추출
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as zf:
            text_parts = []

            # Find all section files
            section_files = sorted([
                name for name in zf.namelist()
                if name.startswith('Contents/section') and name.endswith('.xml')
            ])

            for section_file in section_files:
                try:
                    content = zf.read(section_file).decode('utf-8')
                    # Parse XML and extract text
                    root = ET.fromstring(content)

                    # Extract all text content from XML
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text.strip())
                        if elem.tail and elem.tail.strip():
                            text_parts.append(elem.tail.strip())
                except Exception:
                    continue

            if text_parts:
                return '\n'.join(text_parts)
            else:
                return "HWPX: No text content found in sections."

    except zipfile.BadZipFile:
        return "HWPX: Invalid file format."
    except Exception as e:
        return f"HWPX extraction failed: {str(e)}"


def extract_from_pdf(file_bytes: bytes) -> str:
    """
    PDF 파일에서 페이지별 텍스트 추출
    """
    if not HAS_PYPDF:
        return "PDF extraction requires pypdf library."

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {i+1}]\n{page_text}")
            except Exception:
                continue

        if text_parts:
            return '\n\n'.join(text_parts)
        else:
            return "PDF: No extractable text found (may be image-based)."

    except Exception as e:
        return f"PDF extraction failed: {str(e)}"


def extract_from_docx(file_bytes: bytes) -> str:
    """
    DOCX 파일에서 문단별 텍스트 추출
    """
    if not HAS_DOCX:
        return "DOCX extraction requires python-docx library."

    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        if text_parts:
            return '\n'.join(text_parts)
        else:
            return "DOCX: No text content found."

    except Exception as e:
        return f"DOCX extraction failed: {str(e)}"


def extract_from_xlsx(file_bytes: bytes) -> str:
    """
    XLSX 파일에서 셀 값 추출
    """
    if not HAS_OPENPYXL:
        return "XLSX extraction requires openpyxl library."

    try:
        wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f"[Sheet: {sheet_name}]"]

            for row in sheet.iter_rows(max_row=500):  # Limit rows
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    sheet_text.append(' | '.join(row_values))

            if len(sheet_text) > 1:
                text_parts.append('\n'.join(sheet_text))

        wb.close()

        if text_parts:
            return '\n\n'.join(text_parts)
        else:
            return "XLSX: No data found."

    except Exception as e:
        return f"XLSX extraction failed: {str(e)}"


def extract_from_zip(file_bytes: bytes, original_url: str = "") -> str:
    """
    ZIP 파일 처리 - 우선순위에 따라 최적 파일 선택 후 추출
    """
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), 'r') as zf:
            file_list = zf.namelist()

            best_file = select_best_file_from_zip(file_list)

            if not best_file:
                return f"ZIP: No suitable document found. Files in archive: {', '.join(file_list[:10])}"

            # Extract the best file
            inner_bytes = zf.read(best_file)
            inner_ext = Path(best_file).suffix.lower()

            # Recursively process based on extension
            if inner_ext == '.hwp':
                result = extract_from_hwp(inner_bytes)
            elif inner_ext == '.hwpx':
                result = extract_from_hwpx(inner_bytes)
            elif inner_ext == '.pdf':
                result = extract_from_pdf(inner_bytes)
            elif inner_ext == '.docx':
                result = extract_from_docx(inner_bytes)
            elif inner_ext == '.xlsx':
                result = extract_from_xlsx(inner_bytes)
            else:
                result = f"ZIP: Unsupported inner file format: {inner_ext}"

            return f"[Extracted from ZIP: {best_file}]\n\n{result}"

    except zipfile.BadZipFile:
        return "ZIP: Invalid or corrupted archive."
    except Exception as e:
        return f"ZIP extraction failed: {str(e)}"


def extract_text_from_bytes(file_bytes: bytes, filename: str, url: str = "") -> str:
    """
    바이트 데이터에서 파일 형식에 따라 텍스트 추출 (메인 디스패처)
    """
    ext = Path(filename).suffix.lower()

    if ext == '.zip':
        return extract_from_zip(file_bytes, url)
    elif ext == '.hwp':
        return extract_from_hwp(file_bytes)
    elif ext == '.hwpx':
        return extract_from_hwpx(file_bytes)
    elif ext == '.pdf':
        return extract_from_pdf(file_bytes)
    elif ext == '.docx':
        return extract_from_docx(file_bytes)
    elif ext in ('.xlsx', '.xls'):
        return extract_from_xlsx(file_bytes)
    else:
        return f"Unsupported file format: {ext}. Please check the manual link: {url}"


async def download_file(url: str, timeout: float = 60.0) -> tuple[bytes, str]:
    """
    URL에서 파일 다운로드
    Returns: (file_bytes, filename)
    """
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
        response = await client.get(url)
        response.raise_for_status()

        # Try to get filename from Content-Disposition header
        filename = ""
        content_disposition = response.headers.get('content-disposition', '')
        if 'filename=' in content_disposition:
            # Handle both filename= and filename*=
            import urllib.parse
            if 'filename*=' in content_disposition:
                # RFC 5987 encoded filename
                match = re.search(r"filename\*=(?:UTF-8''|utf-8'')(.+?)(?:;|$)", content_disposition, re.IGNORECASE)
                if match:
                    filename = urllib.parse.unquote(match.group(1))
            if not filename:
                match = re.search(r'filename="?([^";\n]+)"?', content_disposition)
                if match:
                    filename = match.group(1)

        # Fallback to URL path
        if not filename:
            from urllib.parse import urlparse
            filename = os.path.basename(urlparse(url).path)

        return response.content, filename


async def extract_text_from_url(url: str, filename: str = "") -> str:
    """
    URL에서 파일 다운로드 후 텍스트 추출 (메인 진입점)

    Args:
        url: 파일 다운로드 URL
        filename: 파일명 (없으면 URL에서 추출)

    Returns:
        추출된 텍스트 또는 에러 메시지
    """
    try:
        file_bytes, detected_filename = await download_file(url)

        # Use provided filename or detected one
        final_filename = filename if filename else detected_filename

        if not final_filename:
            return f"Text extraction unavailable: Could not determine file type. Manual link: {url}"

        return extract_text_from_bytes(file_bytes, final_filename, url)

    except httpx.HTTPStatusError as e:
        return f"Download failed (HTTP {e.response.status_code}). Manual link: {url}"
    except httpx.TimeoutException:
        return f"Download timed out. Manual link: {url}"
    except Exception as e:
        return f"Text extraction unavailable: {str(e)}. Manual link: {url}"
